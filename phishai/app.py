"""
EndPhishAI Flask API Server with Enhanced Content Analysis
Multi-layer phishing detection with optional deep content scanning
Production-ready with professional error handling
"""

import os
from dotenv import load_dotenv
load_dotenv()
import re
import hashlib
import time
from datetime import datetime, timedelta
from flask import Flask, request, jsonify
from flask_cors import CORS
import requests
from urllib.parse import urlparse
from pathlib import Path
from phishing_detector import detect_phishing

app = Flask(__name__)
CORS(app)

# API Configuration
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY', '')
VIRUSTOTAL_API_KEY = os.getenv('VIRUSTOTAL_API_KEY', '')
GOOGLE_SAFE_BROWSING_URL = "https://safebrowsing.googleapis.com/v4/threatMatches:find"

# Twilio SMS Configuration
TWILIO_ACCOUNT_SID = os.getenv('TWILIO_ACCOUNT_SID', '')
TWILIO_AUTH_TOKEN = os.getenv('TWILIO_AUTH_TOKEN', '')
TWILIO_PHONE_NUMBER = os.getenv('TWILIO_PHONE_NUMBER', '')

# Initialize Twilio client if credentials are available
twilio_client = None
SMS_AVAILABLE = False

# Try to import Twilio - it's optional
try:
    from twilio.rest import Client
    TWILIO_IMPORTED = True
except ImportError:
    TWILIO_IMPORTED = False
    print("Twilio SMS: Package not installed (optional feature)")

# Only initialize if package is available and credentials exist
if TWILIO_IMPORTED and TWILIO_ACCOUNT_SID and TWILIO_AUTH_TOKEN and TWILIO_PHONE_NUMBER:
    try:
        twilio_client = Client(TWILIO_ACCOUNT_SID, TWILIO_AUTH_TOKEN)
        SMS_AVAILABLE = True
        print("‚úÖ Twilio SMS: Configured and ready")
    except Exception as e:
        print("Twilio SMS: Configuration error")
elif TWILIO_IMPORTED:
    print("Twilio SMS: Not configured (add credentials to .env)")
else:
    print("Twilio SMS: Not available (install with: pip install twilio)")

# Try to import python-magic for file validation
try:
    import magic
    MAGIC_AVAILABLE = True
    print("‚úÖ python-magic: Loaded (file security enabled)")
except ImportError:
    MAGIC_AVAILABLE = False
    print("‚ö†Ô∏è python-magic: Not installed (MIME validation disabled)")
    print("   Install with: pip install python-magic-bin (Windows)")
    print("   or: pip install python-magic (Linux/Mac)")


# ============================================================
# SAFETY WRAPPER FOR URL CONTENT ANALYSIS
# ============================================================
class SafeURLAnalyzer:
    """
    Safety wrapper with rate limiting, caching, and blacklist
    """
    
    def __init__(self):
        self.request_times = []
        self.max_requests_per_minute = 10
        self.cache = {}
        self.cache_ttl = timedelta(hours=1)
        
        # Domains to NEVER visit
        self.blacklist = {
            'localhost', '127.0.0.1', '0.0.0.0', '192.168.',
            'internal', 'intranet', 'admin', '10.0.0.'
        }
        
        # Known safe domains (skip analysis)
        self.whitelist = {
            'google.com', 'microsoft.com', 'apple.com', 'facebook.com',
            'amazon.com', 'netflix.com', 'twitter.com', 'linkedin.com',
            'instagram.com', 'youtube.com', 'github.com', 'stackoverflow.com',
            'wikipedia.org', 'reddit.com'
        }
    
    def analyze_safely(self, url):
        """
        Analyze URL with all safety checks
        """
        # Validate URL format
        if not url.startswith(('http://', 'https://')):
            return {
                'success': False,
                'error': 'Invalid URL format',
                'phishing_score': 0.0
            }
        
        # Extract domain
        try:
            domain = urlparse(url).netloc.lower()
        except:
            return {
                'success': False,
                'error': 'Invalid URL',
                'phishing_score': 0.0
            }
        
        # Check blacklist
        if any(blocked in domain for blocked in self.blacklist):
            return {
                'success': False,
                'error': 'URL blocked for safety (internal network)',
                'phishing_score': 0.0
            }
        
        # Check whitelist
        if any(trusted in domain for trusted in self.whitelist):
            return {
                'success': True,
                'url': url,
                'phishing_score': 0.0,
                'is_phishing': False,
                'indicators_found': ['Whitelisted domain - analysis skipped'],
                'analysis': {'total_score': 0.0, 'indicators': []},
                'whitelisted': True
            }
        
        # Check cache
        if url in self.cache:
            cached_result, cached_time = self.cache[url]
            if datetime.now() - cached_time < self.cache_ttl:
                cached_result['cached'] = True
                return cached_result
        
        # Rate limiting
        now = time.time()
        self.request_times = [t for t in self.request_times if now - t < 60]
        
        if len(self.request_times) >= self.max_requests_per_minute:
            return {
                'success': False,
                'error': 'Rate limit exceeded. Please wait a moment.',
                'phishing_score': 0.0
            }
        
        # Perform analysis
        self.request_times.append(now)
        
        try:
            from url_content_analyzer import URLContentAnalyzer
            analyzer = URLContentAnalyzer()
            result = analyzer.analyze_url_content(url)
            
            # Cache successful results
            if result.get('success'):
                self.cache[url] = (result, datetime.now())
            
            return result
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Analysis error: {str(e)}',
                'phishing_score': 0.0
            }


# Initialize safe analyzer
safe_analyzer = SafeURLAnalyzer()
# ============================================================


def sanitize_input(text):
    """Remove potentially harmful characters from user input"""
    if not text or not isinstance(text, str):
        return ""
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Remove control characters
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    # Normalize whitespace
    text = ' '.join(text.split())
    return text.strip()[:5000]


def check_google_safe_browsing(url):
    """Check URL against Google Safe Browsing database"""
    if not GOOGLE_API_KEY:
        return {'result': 'unknown', 'details': 'Google API key not configured'}
    
    try:
        payload = {
            "client": {"clientId": "phishai", "clientVersion": "1.0.0"},
            "threatInfo": {
                "threatTypes": ["MALWARE", "SOCIAL_ENGINEERING", "UNWANTED_SOFTWARE", "POTENTIALLY_HARMFUL_APPLICATION"],
                "platformTypes": ["ANY_PLATFORM"],
                "threatEntryTypes": ["URL"],
                "threatEntries": [{"url": url}]
            }
        }
        
        response = requests.post(f"{GOOGLE_SAFE_BROWSING_URL}?key={GOOGLE_API_KEY}", json=payload, timeout=8)
        
        if response.status_code == 200:
            data = response.json()
            if 'matches' in data and data['matches']:
                threat_types = [match['threatType'] for match in data['matches']]
                return {
                    'result': 'malicious',
                    'details': {
                        'threats': threat_types,
                        'message': 'URL flagged by Google Safe Browsing'
                    }
                }
            return {'result': 'safe', 'details': 'No threats detected by Google Safe Browsing'}
        else:
            return {'result': 'error', 'details': f'API returned status {response.status_code}'}
    
    except requests.Timeout:
        return {'result': 'error', 'details': 'Request timeout'}
    except Exception as e:
        return {'result': 'error', 'details': f'Error: {str(e)}'}


def check_phishtank(url):
    """Check URL against PhishTank community database"""
    try:
        phishtank_url = "http://checkurl.phishtank.com/checkurl/"
        payload = {'url': url, 'format': 'json'}
        response = requests.post(phishtank_url, data=payload, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            if 'results' in data and data['results'].get('in_database'):
                if data['results'].get('valid'):
                    return {
                        'result': 'malicious',
                        'details': {
                            'source': 'PhishTank',
                            'verified': True,
                            'message': 'URL verified as phishing by PhishTank community'
                        }
                    }
            return {'result': 'safe', 'details': 'Not found in PhishTank database'}
        else:
            return {'result': 'unknown', 'details': 'PhishTank check failed'}
    
    except requests.Timeout:
        return {'result': 'error', 'details': 'PhishTank timeout'}
    except Exception as e:
        return {'result': 'error', 'details': f'PhishTank error: {str(e)}'}


def check_urlhaus(url):
    """Check URL against URLhaus malware database"""
    try:
        urlhaus_api = "https://urlhaus-api.abuse.ch/v1/url/"
        payload = {'url': url}
        response = requests.post(urlhaus_api, data=payload, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('query_status') == 'ok':
                threat = data.get('threat', 'unknown')
                return {
                    'result': 'malicious',
                    'details': {
                        'source': 'URLhaus',
                        'threat_type': threat,
                        'message': f'URL flagged as {threat} by URLhaus'
                    }
                }
            else:
                return {'result': 'safe', 'details': 'Not found in URLhaus database'}
        else:
            return {'result': 'unknown', 'details': 'URLhaus check failed'}
    
    except requests.Timeout:
        return {'result': 'error', 'details': 'URLhaus timeout'}
    except Exception as e:
        return {'result': 'error', 'details': f'URLhaus error: {str(e)}'}


def check_virustotal(url):
    """Check URL against VirusTotal's multi-engine scanner"""
    if not VIRUSTOTAL_API_KEY:
        return {'result': 'unknown', 'details': 'VirusTotal API key not configured'}
    
    try:
        headers = {
            'x-apikey': VIRUSTOTAL_API_KEY,
            'User-Agent': 'PhishAI-Educational-Project/1.0'
        }
        
        # VirusTotal uses SHA256 hash of URL as identifier
        url_id = hashlib.sha256(url.encode()).hexdigest()
        report_url = f"https://www.virustotal.com/api/v3/urls/{url_id}"
        response = requests.get(report_url, headers=headers, timeout=8)
        
        if response.status_code == 200:
            data = response.json()
            stats = data['data']['attributes']['last_analysis_stats']
            
            malicious = stats.get('malicious', 0)
            suspicious = stats.get('suspicious', 0)
            total_engines = sum(stats.values())
            
            if malicious >= 2:
                return {
                    'result': 'malicious',
                    'details': {
                        'source': 'VirusTotal',
                        'malicious_engines': malicious,
                        'suspicious_engines': suspicious,
                        'total_engines': total_engines,
                        'message': f'Detected by {malicious} antivirus engines'
                    }
                }
            elif malicious > 0 or suspicious > 0:
                return {
                    'result': 'suspicious',
                    'details': {
                        'source': 'VirusTotal',
                        'malicious_engines': malicious,
                        'suspicious_engines': suspicious,
                        'total_engines': total_engines,
                        'message': f'Flagged by {malicious + suspicious} engines'
                    }
                }
            else:
                return {
                    'result': 'safe',
                    'details': {
                        'source': 'VirusTotal',
                        'malicious_engines': 0,
                        'suspicious_engines': 0,
                        'total_engines': total_engines,
                        'message': 'No threats detected by VirusTotal'
                    }
                }
        
        elif response.status_code == 404:
            return {
                'result': 'unknown',
                'details': {
                    'source': 'VirusTotal',
                    'message': 'URL not in database (normal for new URLs)'
                }
            }
        else:
            return {'result': 'error', 'details': f'VirusTotal API error: {response.status_code}'}
    
    except requests.Timeout:
        return {'result': 'error', 'details': 'VirusTotal timeout'}
    except Exception as e:
        return {'result': 'error', 'details': f'VirusTotal error: {str(e)}'}


def extract_host(url_or_text):
    """Extract hostname from URL"""
    try:
        if url_or_text.startswith('http'):
            parsed = urlparse(url_or_text)
            return parsed.netloc or 'unknown'
        return 'N/A (not a URL)'
    except:
        return 'unknown'


def is_dangerous_file_type(filename):
    """Check if file extension is potentially dangerous"""
    dangerous_extensions = {
        '.exe', '.scr', '.bat', '.cmd', '.com', '.pif', '.application',
        '.gadget', '.msi', '.msp', '.hta', '.cpl', '.jar', '.js', '.jse',
        '.vb', '.vbe', '.ws', '.wsf', '.wsc', '.wsh', '.ps1', '.ps1xml',
        '.ps2', '.ps2xml', '.psc1', '.psc2', '.msh', '.msh1', '.msh2',
        '.mshxml', '.msh1xml', '.msh2xml', '.scf', '.lnk', '.inf', '.reg'
    }
    
    file_ext = '.' + filename.lower().split('.')[-1] if '.' in filename else ''
    return file_ext in dangerous_extensions, file_ext


# ============================================================
# SECURE FILE EXTRACTION FUNCTIONS
# ============================================================

def extract_text_from_file(file, filename):
    """
    SECURE text extraction with multiple safety layers
    Prevents: Code injection, path traversal, memory attacks
    """
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit
    MAX_TEXT_LENGTH = 500000  # 500KB text limit
    
    # ===== SECURITY LAYER 1: File Size Check =====
    file.seek(0, 2)
    file_size = file.tell()
    file.seek(0)
    
    if file_size > MAX_FILE_SIZE:
        return f"[ERROR: File too large - {file_size / (1024*1024):.1f}MB exceeds 10MB limit]"
    
    # ===== SECURITY LAYER 2: Extension Validation =====
    file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
    
    ALLOWED_EXTENSIONS = {'pdf', 'txt', 'csv', 'html', 'htm', 'docx', 'md'}
    DANGEROUS_EXTENSIONS = {
        'exe', 'scr', 'bat', 'cmd', 'com', 'pif', 'application', 'msi',
        'js', 'jse', 'vbs', 'vbe', 'ps1', 'sh', 'dll', 'sys', 'jar'
    }
    
    if file_ext in DANGEROUS_EXTENSIONS:
        return f"[üö´ BLOCKED: {file_ext.upper()} files are prohibited for security]"
    
    if file_ext not in ALLOWED_EXTENSIONS:
        return f"[ERROR: Unsupported file type .{file_ext}]"
    
    # ===== SECURITY LAYER 3: MIME Type Validation (if available) =====
    if MAGIC_AVAILABLE:
        try:
            mime = magic.from_buffer(file.read(2048), mime=True)
            file.seek(0)
            
            ALLOWED_MIMES = {
                'application/pdf',
                'text/plain',
                'text/csv',
                'text/html',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'text/markdown'
            }
            
            if mime not in ALLOWED_MIMES:
                return f"[üö´ BLOCKED: File type mismatch - claimed .{file_ext}, actually {mime}]"
        
        except Exception as e:
            return f"[ERROR: File validation failed - {str(e)[:50]}]"
    
    # ===== SECURITY LAYER 4: Filename Sanitization =====
    safe_filename = os.path.basename(filename)
    if '..' in safe_filename or '/' in safe_filename or '\\' in safe_filename:
        return "[üö´ BLOCKED: Invalid filename - path traversal attempt detected]"
    
    # ===== EXTRACTION BY FILE TYPE =====
    extracted_text = ""
    
    try:
        if file_ext == 'pdf':
            extracted_text = _extract_pdf_secure(file)
        
        elif file_ext == 'docx':
            extracted_text = _extract_docx_secure(file)
        
        elif file_ext in ['txt', 'csv', 'md']:
            extracted_text = _extract_text_secure(file)
        
        elif file_ext in ['html', 'htm']:
            extracted_text = _extract_html_secure(file)
        
        else:
            return f"[ERROR: Handler not implemented for .{file_ext}]"
        
        # ===== SECURITY LAYER 5: Output Sanitization =====
        if len(extracted_text) > MAX_TEXT_LENGTH:
            extracted_text = extracted_text[:MAX_TEXT_LENGTH] + "\n[...truncated for security]"
        
        # Remove potential XSS/injection characters
        extracted_text = _sanitize_output(extracted_text)
        
        file.seek(0)
        return extracted_text.strip()
    
    except Exception as e:
        return f"[ERROR: Extraction failed - {str(e)[:100]}]"


def _extract_pdf_secure(file):
    """Secure PDF extraction"""
    try:
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(file)
        
        MAX_PAGES = 50
        page_count = min(len(pdf_reader.pages), MAX_PAGES)
        
        text = ""
        for i in range(page_count):
            try:
                page_text = pdf_reader.pages[i].extract_text()
                if page_text:
                    text += page_text + "\n"
            except:
                continue
        
        if not text.strip():
            return "[PDF appears to be image-based or encrypted]"
        
        return text
    
    except ImportError:
        return "[ERROR: PyPDF2 not installed - run: pip install PyPDF2]"
    except Exception as e:
        return f"[PDF error: {str(e)[:50]}]"


def _extract_docx_secure(file):
    """Secure DOCX extraction"""
    try:
        import docx
        doc = docx.Document(file)
        
        text = ""
        MAX_PARAGRAPHS = 500
        
        for i, para in enumerate(doc.paragraphs):
            if i >= MAX_PARAGRAPHS:
                text += "\n[...truncated...]"
                break
            text += para.text + "\n"
        
        for table in doc.tables[:10]:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            return "[DOCX appears empty]"
        
        return text
    
    except ImportError:
        return "[ERROR: python-docx not installed - run: pip install python-docx]"
    except Exception as e:
        return f"[DOCX error: {str(e)[:50]}]"


def _extract_text_secure(file):
    """Secure plain text extraction"""
    try:
        text = file.read().decode('utf-8', errors='ignore')
        
        if not text.strip():
            return "[File appears empty]"
        
        return text
    
    except Exception as e:
        return f"[Text error: {str(e)[:50]}]"


def _extract_html_secure(file):
    """Secure HTML extraction"""
    try:
        from bs4 import BeautifulSoup
        
        html_content = file.read().decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for script in soup(['script', 'style', 'meta', 'link']):
            script.decompose()
        
        text = soup.get_text(separator=' ', strip=True)
        
        if not text.strip():
            return "[HTML contains no text]"
        
        return f"HTML Content:\n{text}"
    
    except ImportError:
        return "[ERROR: beautifulsoup4 not installed - run: pip install beautifulsoup4]"
    except Exception as e:
        return f"[HTML error: {str(e)[:50]}]"


def _sanitize_output(text):
    """Remove dangerous characters"""
    if not text:
        return ""
    
    text = text.replace('\x00', '')
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    text = re.sub(r' {10,}', '  ', text)
    
    return text


@app.route("/")
def home():
    """Health check endpoint"""
    return jsonify({
        "status": "running",
        "service": "PhishAI Detection Service",
        "version": "5.0",
        "ml_available": True,
        "file_upload_available": True,
        "sms_available": SMS_AVAILABLE,
        "content_analysis_available": True,
        "google_api_configured": bool(GOOGLE_API_KEY),
        "virustotal_api_configured": bool(VIRUSTOTAL_API_KEY),
        "protection_layers": 7,
        "databases": [
            "Google Safe Browsing", 
            "VirusTotal", 
            "PhishTank",
            "URLhaus",
            "Custom ML Model",
            "File Upload Scanner",
            "Content Analysis (with safety wrapper)"
        ],
        "supported_files": ["PDF", "TXT", "CSV", "HTML", "DOCX"]
    })


@app.route("/recent-threats", methods=["GET"])
def recent_threats():
    """Get recent threat statistics"""
    recent_threats_data = {
        "recent_threats": [
            {"type": "African Banking", "description": "Fake GTBank login pages", "count": 47, "region": "West Africa"},
            {"type": "Mobile Money", "description": "M-Pesa impersonation scams", "count": 23, "region": "East Africa"},
            {"type": "Global Brands", "description": "Fake Facebook/Netflix pages", "count": 156, "region": "Global"},
            {"type": "Payment Processors", "description": "Fake PayPal verification", "count": 89, "region": "Global"}
        ],
        "total_detected": 315,
        "last_updated": "2024-01-15 14:30:00"
    }
    return jsonify(recent_threats_data)


@app.route("/predict", methods=["POST"])
def predict():
    """
    Main prediction endpoint for URL/text analysis
    NOW USES SAFE ANALYZER with rate limiting and caching
    """
    try:
        data = request.get_json(force=True)
        input_text = data.get('url') or data.get('text') or ""
        use_content_analysis = data.get('content_analysis', False)
        
        if not input_text:
            return jsonify({"error": "missing_input", "message": "Please provide 'url' or 'text' field"}), 400
        
        # Sanitize input to prevent injection attacks
        input_text = sanitize_input(input_text)
        
        if not input_text:
            return jsonify({"error": "invalid_input", "message": "Input contained only invalid characters"}), 400
        
        # Validate content analysis parameter
        if not isinstance(use_content_analysis, bool):
            use_content_analysis = False
        
        # Extract hostname for display
        host = extract_host(input_text)
        
        # Run AI/ML phishing detection with optional content analysis
        # THIS WILL USE THE SAFE ANALYZER
        ai_result = detect_phishing(input_text, use_content_analysis=use_content_analysis, safe_analyzer=safe_analyzer)
        
        # Initialize external API results
        google_result = {'result': 'unknown', 'details': 'Not a URL'}
        virustotal_result = {'result': 'unknown', 'details': 'Not a URL'}
        phishtank_result = {'result': 'unknown', 'details': 'Not a URL'}
        urlhaus_result = {'result': 'unknown', 'details': 'Not a URL'}
        
        # Only run URL checks if input is a valid URL
        if input_text.startswith('http'):
            google_result = check_google_safe_browsing(input_text)
            virustotal_result = check_virustotal(input_text)
            phishtank_result = check_phishtank(input_text)
            urlhaus_result = check_urlhaus(input_text)
        
        # Aggregate results from all threat databases
        threat_sources = []
        
        if google_result['result'] == 'malicious':
            threat_sources.append('Google Safe Browsing')
        if virustotal_result['result'] == 'malicious':
            threat_sources.append('VirusTotal')
        if phishtank_result['result'] == 'malicious':
            threat_sources.append('PhishTank')
        if urlhaus_result['result'] == 'malicious':
            threat_sources.append('URLhaus')
        
        # Determine final status based on all checks
        if threat_sources:
            final_status = 'phishing'
            final_explanation = f"BLOCKED by {', '.join(threat_sources)}. {ai_result['explanation']}"
        elif ai_result['prediction'] == 'phishing':
            final_status = 'phishing'
            final_explanation = ai_result['explanation']
        elif ai_result['prediction'] == 'suspicious':
            final_status = 'suspicious'
            final_explanation = ai_result['explanation']
        else:
            final_status = 'safe'
            final_explanation = ai_result['explanation']
        
        # Convert numpy floats to Python floats for JSON serialization
        def convert_floats(obj):
            if isinstance(obj, float):
                return float(obj)
            elif isinstance(obj, dict):
                return {k: convert_floats(v) for k, v in obj.items()}
            elif isinstance(obj, list):
                return [convert_floats(v) for v in obj]
            return obj
        
        # Build comprehensive response
        response = {
            "url": input_text,
            "host": host,
            "ai_result": ai_result['prediction'],
            "ai_score": float(ai_result['confidence']) if ai_result['confidence'] is not None else 0.0,
            "ai_reason": ai_result['explanation'],
            "google_result": google_result['result'],
            "google_details": google_result['details'],
            "virustotal_result": virustotal_result['result'],
            "virustotal_details": virustotal_result['details'],
            "phishtank_result": phishtank_result['result'],
            "urlhaus_result": urlhaus_result['result'],
            "threat_sources": threat_sources,
            "final_status": final_status,
            "explanation": final_explanation,
            "method": ai_result.get('method', 'Heuristic Analysis'),
            "heuristic_score": float(ai_result.get('heuristic_score', 0.0)) if ai_result.get('heuristic_score') is not None else 0.0,
            "ml_score": float(ai_result.get('ml_score', 0.0)) if ai_result.get('ml_score') is not None else None,
            "content_analysis_performed": ai_result.get('content_analysis_performed', False),
            "content_analysis_available": True,
            "protection_layers": 6
        }
        
        # =====  ADD NLP ANALYSIS RESULTS =====
        if 'nlp_analysis' in ai_result:
            response['nlp_analysis'] = ai_result['nlp_analysis']
        # =========================================
        
        
       # Add content analysis details if performed
        if ai_result.get('content_analysis_performed'):
            response["content_indicators"] = ai_result.get('content_indicators', [])
        elif ai_result.get('content_analysis_error'):
            response["content_analysis_error"] = ai_result.get('content_analysis_error')
        
        # ===== ADD THIS NEW CODE HERE =====
        # Extract content analysis stats from ai_result
       # ===== FIXED: Extract content analysis stats =====
        if ai_result.get('content_analysis_performed'):
            # PRIORITY 1: Check content_analysis_data (set by phishing_detector)
            if 'content_analysis_data' in ai_result:
                content_data = ai_result['content_analysis_data']
                print(f"‚úÖ Found content_analysis_data: {content_data}")
            # PRIORITY 2: Check analysis field  
            elif 'analysis' in ai_result and isinstance(ai_result['analysis'], dict):
                analysis = ai_result['analysis']
                content_data = {
                    'html_elements': analysis.get('html_elements', 0),
                    'scripts_count': analysis.get('scripts_count', 0),
                    'forms_count': analysis.get('forms_count', 0),
                    'external_links': analysis.get('external_links', 0)
                }
                print(f"‚úÖ Extracted from analysis field: {content_data}")
            else:
                content_data = {
                    'html_elements': 0,
                    'scripts_count': 0,
                    'forms_count': 0,
                    'external_links': 0
                }
                print(f"‚ö†Ô∏è No content data found, using zeros")
            
            # Add to response
            response["content_analysis"] = content_data
            
            # Add content_score
            if 'phishing_score' in ai_result:
                response["content_score"] = ai_result['phishing_score']
            else:
                response["content_score"] = ai_result.get('confidence', 0.0)
            
            # Add the indicators
            response["content_indicators"] = ai_result.get('content_indicators', [])
            
            print(f"üì§ Sending to frontend: content_analysis={content_data}, content_score={response['content_score']}")
        elif ai_result.get('content_analysis_error'):
            response["content_analysis_error"] = ai_result.get('content_analysis_error')
        # ===== END FIXED SECTION =====
        
        response = convert_floats(response)
        return jsonify(response), 200
    
    except Exception as e:
        print(f"Error in /predict: {e}")
        return jsonify({"error": "internal_error", "message": "Analysis failed", "details": str(e)}), 500


@app.route("/upload-file", methods=["POST"])
def upload_file():
    """
    File upload endpoint for document analysis
    Supports PDF, TXT, CSV, HTML, DOCX files
    """
    try:
        if 'file' not in request.files:
            return jsonify({"error": "no_file", "message": "Please upload a file"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"error": "no_file", "message": "No file selected"}), 400
        
        # Check for dangerous file types
        is_dangerous, file_ext = is_dangerous_file_type(file.filename)
        if is_dangerous:
            return jsonify({
                "error": "dangerous_file",
                "message": f"DANGEROUS FILE TYPE: {file_ext} files can contain malware"
            }), 400
        
        # Validate file extension
        allowed_extensions = {'pdf', 'txt', 'csv', 'html', 'docx'}
        file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
        
        if file_extension not in allowed_extensions:
            return jsonify({
                "error": "invalid_file_type", 
                "message": f"File type .{file_extension} not supported. Allowed: {', '.join(allowed_extensions)}"
            }), 400
        
        # Check file size (5MB limit)
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > 5 * 1024 * 1024:
            return jsonify({
                "error": "file_too_large",
                "message": "File size exceeds 5MB limit"
            }), 400
        
        print(f"Processing file: {file.filename} ({file_size} bytes)")
        
        # Extract text content from file
        extracted_text = extract_text_from_file(file, file.filename)
        
        if not extracted_text or extracted_text.strip() == "":
            return jsonify({
                "error": "no_content", 
                "message": "No readable text found in the file"
            }), 400
        
        print(f"Extracted {len(extracted_text)} characters from file")
        
        # Run AI analysis on extracted text
        ai_result = detect_phishing(extracted_text)
        
        # Extract and analyze URLs found in the file
        urls_found = re.findall(
            r'https?://(?:[a-zA-Z0-9-]+\.)+[a-zA-Z]{2,}(?:/[^\s]*)?',
            extracted_text
        )
        unique_urls = list(set(urls_found))[:10]
        
        # Check first 3 URLs against Google Safe Browsing
        dangerous_urls = []
        for url in unique_urls[:3]:
            try:
                if url.startswith('http'):
                    google_result = check_google_safe_browsing(url)
                    if google_result['result'] == 'malicious':
                        dangerous_urls.append({
                            'url': url,
                            'threat': 'Google Safe Browsing',
                            'details': google_result['details']
                        })
            except:
                pass
        
        # Build file analysis response
        response = {
            "input_type": "file",
            "file_name": file.filename,
            "file_type": file_extension.upper(),
            "file_size": file_size,
            "extracted_text_length": len(extracted_text),
            "extracted_text_preview": extracted_text[:300] + "..." if len(extracted_text) > 300 else extracted_text,
            "urls_found": unique_urls,
            "urls_count": len(unique_urls),
            "dangerous_urls": dangerous_urls,
            "ai_result": ai_result['prediction'],
            "ai_score": float(ai_result['confidence']) if ai_result['confidence'] is not None else 0.0,
            "ai_reason": ai_result['explanation'],
            "final_status": ai_result['prediction'],
            "explanation": f"File Analysis ({file_extension.upper()}): {ai_result['explanation']}",
            "method": f"File Analysis + {ai_result.get('method', 'AI Analysis')}",
            "protection_layers": 6,
            "dangerous_file_blocked": False
        }
        
        # ===== FIX: ADD NLP ANALYSIS FOR FILES TOO =====
        if 'nlp_analysis' in ai_result:
            response['nlp_analysis'] = ai_result['nlp_analysis']
        # =================================================
        
        print(f"File analysis complete: {ai_result['prediction']}")
        return jsonify(response), 200
    
    except Exception as e:
        print(f"Error in /upload-file: {e}")
        return jsonify({
            "error": "internal_error", 
            "message": "File analysis failed", 
            "details": str(e)
        }), 500


@app.route("/send-sms", methods=["POST"])
def send_sms():
    """
    SMS Alert endpoint
    Sends threat notifications via Twilio SMS service
    """
    try:
        if not SMS_AVAILABLE:
            return jsonify({
                "error": "sms_unavailable",
                "message": "SMS service is temporarily unavailable. Please try again later or contact support."
            }), 503
        
        data = request.get_json(force=True)
        phone_number = data.get('phoneNumber')
        threat = data.get('threat', 'unknown')
        url = data.get('url', 'Unknown URL')
        
        if not phone_number:
            return jsonify({
                "error": "missing_phone",
                "message": "Phone number is required"
            }), 400
        
        clean_phone = re.sub(r'\D', '', phone_number)
        if len(clean_phone) < 10:
            return jsonify({
                "error": "invalid_phone",
                "message": "Invalid phone number format. Please include country code (e.g., +234 for Nigeria)"
            }), 400
        
        if not phone_number.startswith('+'):
            formatted_phone = f'+{clean_phone}'
        else:
            formatted_phone = phone_number
        
        if threat in ['phishing', 'malicious']:
            message = (
                "üö® PHISHING ALERT!\n\n"
                f"EndPhishAI detected a DANGEROUS link:\n{url}\n\n"
                "DO NOT CLICK OR SHARE\n\n"
                "Stay safe!\n"
                "- EndPhishAI Security"
            )
        elif threat == 'suspicious':
            message = (
                "‚ö†Ô∏è SUSPICIOUS CONTENT\n\n"
                f"EndPhishAI found suspicious patterns in:\n{url}\n\n"
                "Proceed with caution\n\n"
                "- EndPhishAI Security"
            )
        else:
            message = (
                f"Content verified as safe by PhishAI.\n\n"
                f"URL: {url}\n\n"
                "Stay vigilant!\n"
                "- EndPhishAI Security"
            )
        
        try:
            twilio_message = twilio_client.messages.create(
                body=message,
                from_=TWILIO_PHONE_NUMBER,
                to=formatted_phone
            )
            
            print(f"SMS sent to {formatted_phone}: {twilio_message.sid}")
            
            return jsonify({
                "success": True,
                "message": "SMS alert sent successfully",
                "sid": twilio_message.sid,
                "to": formatted_phone
            }), 200
            
        except Exception as twilio_error:
            error_code = getattr(twilio_error, 'code', None)
            
            if error_code == 21211:
                return jsonify({
                    "error": "invalid_phone",
                    "message": "Invalid phone number. Please check the number and country code."
                }), 400
            elif error_code == 21608:
                return jsonify({
                    "error": "unverified_number",
                    "message": "This phone number needs to be verified. Please contact support."
                }), 400
            elif error_code == 21614:
                return jsonify({
                    "error": "invalid_number",
                    "message": "This phone number cannot receive SMS. Please use a different number."
                }), 400
            else:
                print(f"SMS Error: {twilio_error}")
                return jsonify({
                    "error": "sms_failed",
                    "message": "Unable to send SMS at this time. Please try again later."
                }), 500
    
    except Exception as e:
        print(f"SMS sending error: {e}")
        return jsonify({
            "error": "service_error",
            "message": "Service temporarily unavailable. Please try again later."
        }), 500


if __name__ == "__main__":
    # Startup information
    print("\n" + "="*60)
    print("üöÄ PhishAI Detection Service Starting...")
    print("="*60)
    print(f"Server: http://127.0.0.1:8000")
    print(f"ML Model: Loaded")
    print(f"File Upload: Available")
    print(f"Content Analysis: Available (WITH SAFETY WRAPPER)")
    print(f"SMS Alerts: {'Configured' if SMS_AVAILABLE else 'Not configured (optional)'}")
    print(f"Google API: {'Configured' if GOOGLE_API_KEY else 'Not configured (optional)'}")
    print(f"VirusTotal: {'Configured' if VIRUSTOTAL_API_KEY else 'Not configured (optional)'}")
    print(f"PhishTank: Available")
    print(f"URLhaus: Available")
    print(f"Supported Files: PDF, TXT, CSV, HTML, DOCX")
    print(f"Total Protection Layers: 7")
    print(f"üõ°Ô∏è Safety Features: Rate Limiting, Caching, Blacklist, Whitelist")
    print("="*60 + "\n")
    
    app.run(host="127.0.0.1", port=8000, debug=True)

